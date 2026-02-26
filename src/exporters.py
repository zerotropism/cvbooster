from docx import Document
from markdown import markdown
from bs4 import BeautifulSoup, Tag
from docx.shared import Pt, RGBColor


def _parse_inline_elements(element, paragraph):
    """
    Parse les éléments inline (gras, italique, liens) et les ajoute au paragraphe.

    Args:
        element: L'élément BeautifulSoup à parser
        paragraph: Le paragraphe Word où ajouter le texte
    """
    for content in element.children:
        if isinstance(content, str):
            paragraph.add_run(content)
        elif content.name == "strong" or content.name == "b":
            run = paragraph.add_run(content.get_text())
            run.bold = True
        elif content.name == "em" or content.name == "i":
            run = paragraph.add_run(content.get_text())
            run.italic = True
        elif content.name == "code":
            run = paragraph.add_run(content.get_text())
            run.font.name = "Courier New"
        elif content.name == "a":
            run = paragraph.add_run(content.get_text())
            run.font.color.rgb = RGBColor(0, 0, 255)
            run.underline = True


def to_word(markdown_text):
    """
    Convertit du texte Markdown en document Word (.docx).

    Args:
        markdown_text (str): Le texte au format Markdown à convertir

    Returns:
        Document: Un objet Document python-docx
    """
    # Convertir le markdown en HTML
    html = markdown(markdown_text, extensions=["extra", "nl2br"])

    # Parser le HTML
    soup = BeautifulSoup(html, "html.parser")

    # Créer un nouveau document Word
    doc = Document()

    # Parcourir les éléments HTML et les convertir en éléments Word
    for element in soup.children:
        if not isinstance(element, Tag):
            continue
        if element.name == "h1":
            doc.add_heading(element.get_text(), level=1)
        elif element.name == "h2":
            doc.add_heading(element.get_text(), level=2)
        elif element.name == "h3":
            doc.add_heading(element.get_text(), level=3)
        elif element.name == "p":
            paragraph = doc.add_paragraph()
            _parse_inline_elements(element, paragraph)
        elif element.name == "ul":
            for li in element.find_all("li", recursive=False):
                paragraph = doc.add_paragraph(style="List Bullet")
                _parse_inline_elements(li, paragraph)
        elif element.name == "ol":
            for li in element.find_all("li", recursive=False):
                paragraph = doc.add_paragraph(style="List Number")
                _parse_inline_elements(li, paragraph)
        elif element.name == "pre":
            code = element.get_text()
            doc.add_paragraph(code, style="No Spacing")

    return doc
